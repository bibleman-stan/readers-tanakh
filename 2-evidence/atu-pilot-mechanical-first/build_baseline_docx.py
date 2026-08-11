#!/usr/bin/env python3
"""
build_baseline_docx.py — generate a 3-layer (Hebrew / transliteration / English-gloss)
.docx for Genesis 22 to support Stan's cold-eye ATU baseline.

The output deliberately presents text at VERSE level only, with no pre-applied
ATU segmentation. Stan marks ATU boundaries himself.
"""

from pathlib import Path
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).parent))
import pilot_config as cfg

HE_PATH = cfg.HE_PATH
TRANSLIT_PATH = cfg.TRANSLIT_PATH
ENG_PATH = cfg.ENG_PATH
OUT_PATH = cfg.COLD_EYE_DOCX
VERSE_HEADER_RE = cfg.VERSE_HEADER_RE
CHAPTER_DISPLAY = cfg.CHAPTER_DISPLAY


def parse_verse_file(path: Path) -> dict[int, str]:
    """Return {verse_int: text} from a v0-style file (verse-header line, then content)."""
    verses: dict[int, str] = {}
    current: int | None = None
    buf: list[str] = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            if current is not None and buf:
                verses[current] = " ".join(buf).strip()
                current, buf = None, []
            continue
        m = VERSE_HEADER_RE.match(line)
        if m:
            if current is not None and buf:
                verses[current] = " ".join(buf).strip()
            current = int(m.group(1))
            buf = []
        else:
            buf.append(line)
    if current is not None and buf:
        verses[current] = " ".join(buf).strip()
    return verses


def set_paragraph_rtl(paragraph) -> None:
    """Mark the paragraph as right-to-left for Hebrew rendering in Word."""
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    p_pr.append(bidi)


def set_run_rtl(run) -> None:
    """Mark the run's complex-script (Hebrew) text properties as RTL."""
    r_pr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    r_pr.append(rtl)


def main() -> None:
    he = parse_verse_file(HE_PATH)
    translit_raw = parse_verse_file(TRANSLIT_PATH)
    eng_raw = parse_verse_file(ENG_PATH)

    # Strip ` | ` separators from transliteration so it reads as running text.
    translit = {v: t.replace(" | ", " ").replace("|", " ") for v, t in translit_raw.items()}
    # Collapse multi-line KJV (one ATU per line in v2/eng-kjv) into running prose per verse.
    # parse_verse_file already joined them with spaces; just collapse extra whitespace.
    eng = {v: " ".join(t.split()) for v, t in eng_raw.items()}

    all_verses = sorted(set(he) & set(translit) & set(eng))
    missing_he = sorted(set(he) - set(all_verses))
    missing_translit = sorted(set(translit) - set(all_verses))
    missing_eng = sorted(set(eng) - set(all_verses))
    if missing_he or missing_translit or missing_eng:
        print(f"WARN: verse mismatch — he extras: {missing_he}, translit extras: {missing_translit}, eng extras: {missing_eng}")
    print(f"Verses with all three layers: {len(all_verses)} (range {min(all_verses)}-{max(all_verses)})")

    doc = Document()

    # Page margins — generous left margin for marking
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

    # Title
    title = doc.add_heading(f"{CHAPTER_DISPLAY} — Cold-eye ATU baseline", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Instructions
    instr = doc.add_paragraph()
    instr_run = instr.add_run(
        "Cold-eye protocol: mark ATU boundaries on the HEBREW line by inserting "
        "a forward slash (//) or a line break between adjacent words where you "
        "judge one complete thought ends and another begins. Do NOT consult any "
        "prior rendering, the v2/heb file, the minimal-rubric notes, or the constraint "
        "catalog. Segment from the text alone, using your current judgment of "
        "where an ATU begins and ends."
    )
    instr_run.font.size = Pt(10)
    instr_run.italic = True
    doc.add_paragraph("")

    # Per-verse blocks
    for vnum in all_verses:
        # Verse number heading
        h = doc.add_heading(f"{cfg.VERSE_PREFIX}{vnum}", level=3)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)

        # Hebrew (RTL, larger)
        p_he = doc.add_paragraph()
        p_he.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_rtl(p_he)
        r_he = p_he.add_run(he[vnum])
        # Use "David" as primary; Windows usually substitutes intelligently
        r_he.font.name = "David"
        # Set the complex-script (Hebrew) font name too
        rPr = r_he._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:cs"), "David")
        rFonts.set(qn("w:hAnsi"), "David")
        rPr.append(rFonts)
        r_he.font.size = Pt(18)
        set_run_rtl(r_he)

        # Transliteration (LTR, italic)
        p_tr = doc.add_paragraph()
        r_tr = p_tr.add_run(translit[vnum])
        r_tr.font.size = Pt(11)
        r_tr.italic = True
        r_tr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # English gloss (LTR)
        p_en = doc.add_paragraph()
        r_en = p_en.add_run(eng[vnum])
        r_en.font.size = Pt(11)

        # Marking space
        doc.add_paragraph("")

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote: {OUT_PATH}")


if __name__ == "__main__":
    main()
