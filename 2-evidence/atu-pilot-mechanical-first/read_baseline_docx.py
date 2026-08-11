#!/usr/bin/env python3
"""
read_baseline_docx.py — extract Stan's cold-eye ATU baseline from the
marked-up docx.

The docx structure (per build_baseline_docx.py): for each verse,
  - Heading 3:  "22:N"
  - Paragraph:  Hebrew (RTL)
  - Paragraph:  Transliteration (italic gray)
  - Paragraph:  English KJV running
  - Empty paragraph

Stan reports he marked ATU boundaries on the TRANSLITERATION layer (easier to
read). We extract by walking the document, identifying verse-heading paragraphs,
and reading the transliteration paragraph (or paragraphs, if Stan inserted
paragraph breaks to mark boundaries).

Strategy: for each verse, collect all italic paragraphs between this verse's
heading and the next heading. Each italic paragraph = one ATU (Stan marked
boundaries by inserting paragraph breaks). If Stan used `//` or `|` within
a paragraph instead, we also split on those markers.
"""

from __future__ import annotations
import re
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).parent))
import pilot_config as cfg

DOCX = cfg.COLD_EYE_DOCX
OUT_TXT = cfg.COLD_EYE_EXTRACTED
VERSE_HEADING_RE = cfg.VERSE_HEADER_RE
BOUNDARY_MARKERS_RE = re.compile(r"\s*(?://|\|)\s*")


def main() -> None:
    if not DOCX.exists():
        print(f"ERROR: docx not found: {DOCX}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(DOCX))

    # Walk paragraphs in order; bucket them by verse.
    # A paragraph is the verse-heading if it's a Heading-3 style with "22:N".
    # Italic paragraphs between headings are the transliteration layer.
    current_verse: int | None = None
    verses: dict[int, list[str]] = {}
    italic_paras_per_verse: dict[int, list[str]] = {}

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        m = VERSE_HEADING_RE.match(text)
        if m:
            current_verse = int(m.group(1))
            italic_paras_per_verse.setdefault(current_verse, [])
            continue

        if current_verse is None:
            continue

        # Determine if paragraph is the transliteration layer (italic) or other.
        # The transliteration runs are emitted with run.italic = True.
        is_italic = any(
            (r.italic is True) and r.text.strip() for r in p.runs
        )
        # Fallback: detect by leading/embedded Latin letters with no Hebrew block chars.
        # (If Stan inserted plain-text paragraphs without italic formatting.)
        has_hebrew = any("֐" <= ch <= "׿" for ch in text)
        is_latin_only = not has_hebrew

        # Skip Hebrew and English layers; only take transliteration-like content.
        # Italic OR (Latin-only AND not the English KJV layer).
        # We'll distinguish English vs translit by simple heuristic: translit has
        # diacritic-free transliterated word forms; English has function words.
        # Easiest: just take italic. If italic ran formatting was lost on Stan's edit,
        # take Latin-only paragraphs after the Hebrew but before the English.
        if is_italic:
            italic_paras_per_verse[current_verse].append(text)

    # If italic detection failed because Stan replaced runs (Word can collapse formatting
    # when editing), fall back: per verse, take ALL Latin-only paragraphs minus the last
    # (assumed English) — only if italic count is zero.
    fallback_used: list[int] = []
    if not any(italic_paras_per_verse.values()):
        print("WARN: italic detection found nothing; falling back to Latin-only paragraphs", file=sys.stderr)
        # Re-walk and collect by position
        current_verse = None
        seen_hebrew_for_verse: dict[int, bool] = {}
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            m = VERSE_HEADING_RE.match(text)
            if m:
                current_verse = int(m.group(1))
                italic_paras_per_verse.setdefault(current_verse, [])
                seen_hebrew_for_verse[current_verse] = False
                continue
            if current_verse is None:
                continue
            has_hebrew = any("֐" <= ch <= "׿" for ch in text)
            if has_hebrew:
                seen_hebrew_for_verse[current_verse] = True
                continue
            if seen_hebrew_for_verse.get(current_verse):
                # This is a Latin-only paragraph after the Hebrew. Take it.
                italic_paras_per_verse[current_verse].append(text)
        fallback_used = sorted(italic_paras_per_verse.keys())

    # The fallback would over-include the English line. Trim: drop the LAST
    # Latin-only paragraph per verse (the English KJV running line). The first
    # N-1 are then Stan's transliteration ATUs.
    if fallback_used:
        trimmed: dict[int, list[str]] = {}
        for v, paras in italic_paras_per_verse.items():
            if len(paras) >= 1:
                # Drop the last paragraph (English) — keep the rest as ATU lines
                trimmed[v] = paras[:-1]
            else:
                trimmed[v] = []
        italic_paras_per_verse = trimmed

    # For each verse, expand any in-paragraph markers (// or |) into separate ATU lines.
    final: dict[int, list[str]] = {}
    for v, paras in italic_paras_per_verse.items():
        atus: list[str] = []
        for p in paras:
            # Split on // or | (with optional surrounding whitespace)
            parts = BOUNDARY_MARKERS_RE.split(p)
            atus.extend(part.strip() for part in parts if part.strip())
        final[v] = atus

    # Emit human-readable extraction so Stan can verify
    lines = []
    for v in sorted(final.keys()):
        lines.append(f"22:{v}")
        for atu in final[v]:
            lines.append(f"  {atu}")
        lines.append("")

    OUT_TXT.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote: {OUT_TXT}")

    # Summary
    total_atus = sum(len(v) for v in final.values())
    print(f"\nTotal verses with ATU markings: {len(final)}")
    print(f"Total ATU lines extracted: {total_atus}")
    print(f"\nATUs per verse:")
    for v in sorted(final.keys()):
        print(f"  22:{v}: {len(final[v])} ATU(s)")


if __name__ == "__main__":
    main()
